import os
import sys
from docx.api import Document
import openpyxl
from shutil import copyfile
from tkinter import filedialog
from tkinter import *
from tkinter import ttk
import tkinter as tk

templates = []
forms = []

forms_dir = "asdfghj"
templates_dirs = []

alias_vals = {}
conflicts = {}
labels = []

def check_alias(text):
  for arg in labels:
    if (arg.strip().lower() == text.strip().lower()):
      return arg
  return False


def find_all(a_str, sub):
  start = 0
  while True:
    start = a_str.find(sub, start)
    if start == -1:
      return
    yield start
    start += len(sub)  # use start += 1 to find overlapping matches


def swap_text(paragraph):
  indicies = list(find_all(paragraph.text, '*'))
  if len(indicies) >= 2:
    for label in alias_vals.keys():
      if type(label) is str:
        for i in range(0, len(indicies) - 1):
          if label.lower().strip() == paragraph.text[(indicies[i] + 1):indicies[i + 1]].lower().strip():
            paragraph.text = paragraph.text[:indicies[i]] + \
                alias_vals[label] + paragraph.text[indicies[i + 1]:(len(paragraph.text) - 2)]


def get_files(dir, tab):
  for root, subdirs, files in os.walk(dir):
    if len(files) > 0:
      backup_path = os.path.join(root, 'backup')
      # os.mkdir(backup_path)
      for file in files:
        print(file)
        path = os.path.join(root, file)
        # copyfile(path, backup_path + '/' + file)
        tab.append(path)


def clone_values(doc):
  for table in doc.tables:
    for row_i in range(0, len(table.rows)):
      for col_i, cell in enumerate(table.row_cells(row_i)):
        check = check_alias(cell.text)  # standard table cell check
        if check != False:
          table.cell(row_i, col_i).text = alias_vals[check]
        else:
          for paragraph in cell.paragraphs:  # check for paragraphs in table cell
            swap_text(paragraph)

  for paragraph in doc.paragraphs:
    swap_text(paragraph)

def main():
  print(forms_dir)
  get_files(forms_dir, forms)
  for dir in templates_dirs:
    get_files(dir, templates)

  #pull values
  for form_dir in forms:
    print(form_dir)
    if form_dir.endswith('.docx'):
      doc = Document(form_dir)
      for table in doc.tables:
        for i, row in enumerate(table.rows):
          check = False
          try:
            check = check_alias(row.cells[0].text)
          except IndexError:
            continue

          if check != False:
            cell_text = row.cells[1].text
            if (check in alias_vals.keys()) and (alias_vals[check] != cell_text):
              if (check in conflicts.keys()) == False:
                conflicts[check] = [row.cells[1].text, alias_vals[check]]
              else:
                conflicts[check].append(row.cells[1].text)

            else:
              alias_vals[check] = cell_text

  #resolve conflicting values
  # for label in conflicts:
  #   print("There's conflicting values for label " + label)
  #   for i, value in enumerate(conflicts[label]):
  #     print(str(i) + ") " + value)
  #   sys.stdout.flush()
  #   s = input("Select which value to use: ")
  #   print(s)
  #   sys.stdout.flush()
  #   alias_vals[label] = conflicts[label][int(s)]

  for label in conflicts:
    popup = tk.Toplevel()
    popup.title("Label value conflict")
    text = Label(popup, text=("There's conflicting values for label " + label))
    text.grid(row=0, column=0)
    for i, value in enumerate(conflicts[label]):
      def on_press():
        popup.destroy()
        alias_vals[label] = conflicts[label][i]
      button = Button(popup, text=value, command=on_press)
      button.grid(row=(i + 1), column=0)
    global app
    app.wait_window(popup)

  #clone values
  for template_dir in templates:
    if template_dir.endswith('.docx') and template_dir.find('~') == -1:
      doc = Document(template_dir)
      clone_values(doc)
      for section in doc.sections:
        clone_values(section.header)
        clone_values(section.footer)
      doc.save(template_dir)
    elif template_dir.endswith('.xlsx'):
      wb = openpyxl.load_workbook(template_dir)
      for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        for row in sheet.iter_rows():
          for cell in row:
            cell_text = cell.value
            if (cell_text != 'None' and cell_text is not None):
              for label in alias_vals:
                val = alias_vals[label]
                if type(label) is str:
                  check = list(
                      find_all(cell_text.strip().lower(), label.strip().lower()))
                  if len(check) > 0:
                    # print(cell_text)
                    if len(check) == 1:
                      if (check[0] == 0 and cell_text.strip().lower() == label.strip().lower()):
                        cell.value = val
                      elif check[0] != 0:
                        cell.value = cell_text[:check[0]] + \
                            val + cell_text[(len(val) + check[0]):]
                    else:
                      last_index = check[len(check) - 1]
                      cell.value = cell_text[:last_index] + \
                          val + cell_text[(len(val) + last_index):]
      wb.save(template_dir)




def browse_forms():
  global forms_dir 
  forms_dir = filedialog.askdirectory()
  forms_label_text.set(forms_dir)

def browse_templates():
  new_dir = filedialog.askdirectory()
  templates_dirs.append(new_dir)
  templates_label_text.set('\n'.join(templates_dirs))


def browse_labels():
  file = filedialog.askopenfile(mode='r', filetypes=[('Text Files', '*.txt')])
  if file is not None:
    labels_str = ""
    for line in file:
      labels.append(line)
      labels_str = labels_str + line
    labels_text.set(labels_str)


app = Tk()
app.title("DOS-AI")
app.geometry("600x400")

forms_button = Button(text="Select Forms", command=browse_forms)
forms_button.grid(row=0, column=0)
forms_label_text = StringVar()
forms_label_text.set('forms')
forms_label = Label(app, textvariable=forms_label_text, font=('bold', 14))
forms_label.grid(row=0, column = 1)

templates_button = Button(text="Select Templates", command=browse_templates)
templates_button.grid(row=1, column=0)
templates_label_text = StringVar()
templates_label_text.set('templates')
templates_label = Label(
    app, textvariable=templates_label_text, font=('bold', 14))
templates_label.grid(row=1, column=1)

labels_button = Button(text="Select Labels", command=browse_labels)
labels_button.grid(row=2, column=0)
labels_text = StringVar()
labels_text.set('labels')
labels_label = Label(
    app, textvariable=labels_text, font=('bold', 14))
labels_label.grid(row=2, column=1)

run_button = Button(text="Run", command=main)
run_button.grid(row=3, column=0)

app.mainloop()

